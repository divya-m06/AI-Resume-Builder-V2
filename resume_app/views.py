from django.shortcuts import render, redirect
from django.http import HttpResponse
from django.contrib import messages
from django.views.decorators.csrf import csrf_exempt

from .forms import ResumeForm
from .models import UserAccount

from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from pdfminer.high_level import extract_text

import openai
import spacy
import re


import random
import requests
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.lib import colors
from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from django.http import HttpResponse
from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import json
from django.http import JsonResponse



# ---------- Simple text cleaning for JD matcher ----------
STOP_WORDS = {
    "the", "and", "for", "with", "this", "that", "from", "your", "you",
    "are", "our", "will", "have", "has", "into", "about", "been", "who",
    "able", "such", "their", "them", "they", "any", "all", "etc", "job",
    "role", "responsibilities", "requirements", "skills", "must", "should",
    "experience", "years", "year", "we", "us", "at", "on", "in", "to", "of",
}

def tokenize(text: str):
    text = text.lower()
    # keep letters, numbers, +, #, / (for c++, c#, etc.)
    text = re.sub(r"[^a-z0-9+#/ ]+", " ", text)
    tokens = [t for t in text.split() if t and t not in STOP_WORDS and len(t) > 2]
    return tokens




# ---------------- HOME ----------------
def home(request):
    username = request.session.get("logged_name")
    return render(request, "resume_app/home.html", {
        "username": username
    })



def landing_page(request):
    if request.session.get("logged_user"):
        return redirect("home")
    return render(request, "resume_app/landing.html")

# ---------------- LOGIN ----------------
def login_page(request):
    if request.method == "POST":
        userid = request.POST.get("userid")
        password = request.POST.get("password")

        try:
            user = UserAccount.objects.get(userid=userid, password=password)
            request.session["logged_user"] = user.userid
            request.session["logged_name"] = user.name
            return redirect("home")
        except UserAccount.DoesNotExist:
            messages.error(request, "Invalid UserID or Password")

    return render(request, "resume_app/login.html")


# ---------------- SIGNUP ----------------
def signup_page(request):
    if request.method == "POST":
        name = request.POST.get("name")
        email = request.POST.get("email")
        userid = request.POST.get("userid")
        password = request.POST.get("password")
        phone = request.POST.get("phone")


        if UserAccount.objects.filter(userid=userid).exists():
            messages.error(request, "User ID already exists!")
            return redirect("signup")

        if UserAccount.objects.filter(email=email).exists():
            messages.error(request, "Email already exists!")
            return redirect("signup")

        UserAccount.objects.create(
            name=name,
            email=email,
            phone=phone,
            userid=userid,
            password=password
        )
        messages.success(request, "Account created successfully!")
        return redirect("login")

    return render(request, "resume_app/signup.html")


# ---------------- LOGOUT ----------------
def logout_user(request):
    request.session.flush()
    return redirect("login")



def forgot_password(request):
    if request.method == "POST":
        phone = request.POST.get("phone")

        try:
            user = UserAccount.objects.get(phone=phone)
        except UserAccount.DoesNotExist:
            messages.error(request, "Phone number not registered")
            return redirect("forgot_password")

        otp = random.randint(100000, 999999)

        # Store OTP in session
        request.session["reset_otp"] = otp
        request.session["reset_user"] = user.userid

        # 🔥 Simulated OTP sending
        print(f"OTP for password reset is: {otp}")

        messages.success(request, "OTP sent to registered phone number")
        return redirect("verify_otp")

    return render(request, "resume_app/forgot_password.html")

def reset_password(request):
    if request.method == "POST":
        new_password = request.POST.get("password")
        userid = request.session.get("reset_user")

        user = UserAccount.objects.get(userid=userid)
        user.password = new_password
        user.save()

        # Clear session
        request.session.flush()

        messages.success(request, "Password reset successful")
        return redirect("login")

    return render(request, "resume_app/reset_password.html")

def verify_otp(request):
    if request.method == "POST":
        entered_otp = request.POST.get("otp")

        if str(entered_otp) == str(request.session.get("reset_otp")):
            messages.success(request, "OTP verified")
            return redirect("reset_password")
        else:
            messages.error(request, "Invalid OTP")

    return render(request, "resume_app/verify_otp.html")



# ---------------- SKILL GAP FORM PAGE ----------------
def skill_gap_form(request):
    return render(request, "resume_app/skill_gap.html")



# ---------------- JD ANALYZER FORM PAGE ----------------
def jd_analyzer_form(request):
    return render(request, "resume_app/jd_analyzer.html")



# ---------------- PROCESS JD KEYWORDS ----------------

@csrf_exempt
def jd_analyzer_result(request):

    if request.method != "POST":
        return redirect("jd_analyzer_form")

    # ---------------- INPUTS ----------------
    jd_text = request.POST.get("job_description", "").lower()
    user_projects = request.POST.get("current_projects", "").lower()
    resume_file = request.FILES.get("resume_file")

    resume_text = ""

    # ---------------- EXTRACT RESUME TEXT ----------------
    if resume_file:
        try:
            if resume_file.name.endswith(".pdf"):
                resume_text = extract_text(BytesIO(resume_file.read()))
            elif resume_file.name.endswith(".docx"):
                doc = Document(resume_file)
                resume_text = "\n".join([p.text for p in doc.paragraphs])
        except:
            resume_text = ""

    resume_text = resume_text.lower()

    # ---------------- KEYWORD EXTRACTION ----------------
    jd_keywords = set(re.findall(r'\b[a-zA-Z]{3,}\b', jd_text))
    resume_keywords = set(re.findall(r'\b[a-zA-Z]{3,}\b', resume_text))

    matched_keywords = sorted(jd_keywords & resume_keywords)
    score = int((len(matched_keywords) / len(jd_keywords)) * 100) if jd_keywords else 0

    # ====================================================
    # 🔑 ROLE → SKILL KEYWORD MAP (IMPORTANT)
    # ====================================================
    role_skill_map = {
        "data": ["data", "analysis", "sql", "excel", "dashboard", "statistics"],
        "data scientist": ["python", "machine", "learning", "ml", "deep", "ai"],
        "ai": ["ai", "nlp", "cnn", "deep", "neural", "transformer"],
        "node.js": ["node", "express", "jwt", "api"],
        "full stack": ["frontend", "backend", "react", "node", "api"],
        "django": ["django", "python", "rest", "api"],
        "cloud": ["aws", "cloud", "ec2", "lambda", "s3"],
        "devops": ["docker", "kubernetes", "ci", "cd", "jenkins"],
        "cybersecurity": ["security", "network", "malware", "phishing", "ids"]
    }

    # ====================================================
    # 📚 PROJECT LIBRARY (FULL)
    # ====================================================
    project_library = {

        "data": [
            {
                "title": "Business Sales Dashboard",
                "description": "Interactive dashboard for sales insights.",
                "stack": ["Python", "Pandas", "Power BI"],
                "github": "https://github.com/search?q=sales+dashboard+python",
                "video": "https://www.youtube.com/results?search_query=sales+dashboard+power+bi",
                "article": "https://towardsdatascience.com/data-analysis-dashboard-with-python-8e6d6c0a9e3b"
            },
            {
                "title": "Customer Behavior Analysis",
                "description": "Analyze customer purchasing patterns.",
                "stack": ["Python", "SQL", "EDA"],
                "github": "https://github.com/search?q=customer+behavior+analysis+python",
                "video": "https://www.youtube.com/results?search_query=customer+behavior+analysis",
                "article": "https://towardsdatascience.com/customer-analysis-with-python-6f8a3b9d1e"
            },
            {
                "title": "A/B Testing System",
                "description": "Marketing A/B testing framework.",
                "stack": ["Python", "Statistics"],
                "github": "https://github.com/search?q=ab+testing+python",
                "video": "https://www.youtube.com/results?search_query=ab+testing+python",
                "article": "https://towardsdatascience.com/a-b-testing-with-python-4e5d6a9b2c"
            }
        ],

        "data scientist": [
            {
                "title": "Time Series Forecasting",
                "description": "Forecast trends using ARIMA & LSTM.",
                "stack": ["Python", "Time Series"],
                "github": "https://github.com/search?q=time+series+forecasting+python",
                "video": "https://www.youtube.com/results?search_query=time+series+forecasting",
                "article": "https://towardsdatascience.com/time-series-forecasting-with-python-4c7d1e8f"
            },
            {
                "title": "Fraud Detection System",
                "description": "Detect fraudulent transactions.",
                "stack": ["Python", "ML"],
                "github": "https://github.com/search?q=fraud+detection+machine+learning",
                "video": "https://www.youtube.com/results?search_query=fraud+detection+machine+learning",
                "article": "https://towardsdatascience.com/fraud-detection-with-machine-learning-98e3b1d1f6d"
            },
            {
                "title": "Recommendation Engine",
                "description": "Movie & product recommendations.",
                "stack": ["Python", "ML"],
                "github": "https://github.com/search?q=recommendation+system+python",
                "video": "https://www.youtube.com/results?search_query=recommendation+system+machine+learning",
                "article": "https://towardsdatascience.com/building-a-recommendation-system-with-python-5f4f8c6a1d"
            }
        ],

        "ai": [
            {
                "title": "AI Chatbot",
                "description": "Conversational chatbot using NLP.",
                "stack": ["Python", "NLP"],
                "github": "https://github.com/search?q=ai+chatbot+python",
                "video": "https://www.youtube.com/results?search_query=ai+chatbot+nlp",
                "article": "https://towardsdatascience.com/building-an-ai-chatbot-from-scratch-9f4c3b2a1d"
            },
            {
                "title": "Image Classification",
                "description": "CNN based image classifier.",
                "stack": ["Python", "CNN"],
                "github": "https://github.com/search?q=image+classification+cnn",
                "video": "https://www.youtube.com/results?search_query=image+classification+cnn",
                "article": "https://towardsdatascience.com/image-classification-with-deep-learning-4b7a2e9c1d"
            },
            {
                "title": "Speech Recognition App",
                "description": "Speech to text AI system.",
                "stack": ["Python", "AI"],
                "github": "https://github.com/search?q=speech+recognition+python",
                "video": "https://www.youtube.com/results?search_query=speech+recognition+python",
                "article": "https://towardsdatascience.com/speech-recognition-with-python-1a2b3c4d5e"
            }
        ],

        "node.js": [
            {
                "title": "REST API with Express",
                "description": "Scalable REST API.",
                "stack": ["Node.js", "Express"],
                "github": "https://github.com/search?q=nodejs+rest+api",
                "video": "https://www.youtube.com/results?search_query=nodejs+rest+api",
                "article": "https://www.freecodecamp.org/news/build-a-rest-api-with-node-js-express/"
            },
            {
                "title": "JWT Authentication",
                "description": "Secure auth system.",
                "stack": ["Node.js", "JWT"],
                "github": "https://github.com/search?q=nodejs+jwt+authentication",
                "video": "https://www.youtube.com/results?search_query=nodejs+jwt",
                "article": "https://www.freecodecamp.org/news/how-to-secure-your-node-js-application-with-jwt/"
            },
            {
                "title": "Real-Time Chat App",
                "description": "Socket.io chat application.",
                "stack": ["Node.js", "Socket.io"],
                "github": "https://github.com/search?q=nodejs+chat+application",
                "video": "https://www.youtube.com/results?search_query=nodejs+chat+app",
                "article": "https://www.freecodecamp.org/news/build-a-chat-app/"
            }
        ],

        "full stack": [
            {
                "title": "MERN Job Portal",
                "description": "Full stack job portal.",
                "stack": ["MongoDB", "React", "Node"],
                "github": "https://github.com/search?q=mern+job+portal",
                "video": "https://www.youtube.com/results?search_query=mern+job+portal",
                "article": "https://www.freecodecamp.org/news/build-a-mern-stack-app/"
            },
            {
                "title": "E-Commerce Platform",
                "description": "Full stack ecommerce app.",
                "stack": ["React", "Node"],
                "github": "https://github.com/search?q=full+stack+ecommerce+project",
                "video": "https://www.youtube.com/results?search_query=full+stack+ecommerce",
                "article": "https://www.freecodecamp.org/news/build-an-ecommerce-site/"
            },
            {
                "title": "Social Media Platform",
                "description": "Social networking app.",
                "stack": ["React", "Node"],
                "github": "https://github.com/search?q=full+stack+social+media+project",
                "video": "https://www.youtube.com/results?search_query=full+stack+social+media",
                "article": "https://www.freecodecamp.org/news/build-a-social-network/"
            }
        ],

        "cloud": [
            {
                "title": "Cloud Resume Challenge",
                "description": "Deploy resume on AWS.",
                "stack": ["AWS", "Lambda"],
                "github": "https://github.com/search?q=cloud+resume+challenge",
                "video": "https://www.youtube.com/results?search_query=cloud+resume+challenge",
                "article": "https://cloudresumechallenge.dev/docs/the-challenge/"
            },
            {
                "title": "Serverless API",
                "description": "AWS Lambda API.",
                "stack": ["AWS", "API Gateway"],
                "github": "https://github.com/search?q=aws+serverless+api",
                "video": "https://www.youtube.com/results?search_query=aws+serverless+api",
                "article": "https://www.serverless.com/blog/serverless-rest-api"
            },
            {
                "title": "Cloud Monitoring",
                "description": "AWS CloudWatch monitoring.",
                "stack": ["AWS", "CloudWatch"],
                "github": "https://github.com/search?q=aws+cloudwatch+project",
                "video": "https://www.youtube.com/results?search_query=aws+cloudwatch",
                "article": "https://docs.aws.amazon.com/cloudwatch/"
            }
        ],

        "devops": [
            {
                "title": "CI/CD Pipeline",
                "description": "Jenkins CI/CD setup.",
                "stack": ["Jenkins", "Docker"],
                "github": "https://github.com/search?q=jenkins+cicd+pipeline",
                "video": "https://www.youtube.com/results?search_query=jenkins+cicd",
                "article": "https://www.jenkins.io/doc/tutorials/"
            },
            {
                "title": "Dockerized Microservices",
                "description": "Microservices with Docker.",
                "stack": ["Docker"],
                "github": "https://github.com/search?q=docker+microservices",
                "video": "https://www.youtube.com/results?search_query=docker+microservices",
                "article": "https://www.docker.com/get-started/"
            },
            {
                "title": "Kubernetes Deployment",
                "description": "Deploy apps on Kubernetes.",
                "stack": ["Kubernetes"],
                "github": "https://github.com/search?q=kubernetes+deployment+project",
                "video": "https://www.youtube.com/results?search_query=kubernetes+deployment",
                "article": "https://kubernetes.io/docs/tutorials/"
            }
        ],

        "cybersecurity": [
            {
                "title": "Phishing Detection",
                "description": "Detect phishing URLs.",
                "stack": ["Python", "Security"],
                "github": "https://github.com/search?q=phishing+detection+python",
                "video": "https://www.youtube.com/results?search_query=phishing+detection",
                "article": "https://towardsdatascience.com/phishing-detection-with-machine-learning-3c4d5e6f"
            },
            {
                "title": "Malware Detection",
                "description": "Malware classifier.",
                "stack": ["Python", "ML"],
                "github": "https://github.com/search?q=malware+detection+python",
                "video": "https://www.youtube.com/results?search_query=malware+detection",
                "article": "https://towardsdatascience.com/malware-detection-using-machine-learning-9f4a2d6c"
            },
            {
                "title": "Intrusion Detection",
                "description": "Network IDS system.",
                "stack": ["Python", "IDS"],
                "github": "https://github.com/search?q=intrusion+detection+system",
                "video": "https://www.youtube.com/results?search_query=intrusion+detection",
                "article": "https://towardsdatascience.com/intrusion-detection-system-using-machine-learning-3fdb6d4a9b"
            }
        ]
    }

    # ====================================================
    # 🔥 SMART MATCHING
    # ====================================================
    suggested_projects = []
    for role, skills in role_skill_map.items():
        if any(skill in jd_keywords for skill in skills):
            suggested_projects.extend(project_library.get(role, []))

    # ---------------- FALLBACK ----------------
    if not suggested_projects:
        suggested_projects = [
            {
                "title": "AI Resume Analyzer",
                "description": "AI resume screening project.",
                "stack": ["Python", "NLP"],
                "github": "https://github.com/search?q=resume+analyzer",
                "video": "https://www.youtube.com/results?search_query=resume+analyzer+project",
                "article": "https://medium.com/search?q=resume+analyzer"
            }
        ]

    return render(request, "resume_app/jd_analysis_result.html", {
        "score": score,
        "matched": matched_keywords,
        "suggested_projects": suggested_projects
    })



# ---------------- CREATE RESUME ----------------
def create_resume(request):
    if request.method == "POST":
        form = ResumeForm(request.POST)

        # Extra fields from form (not saved in DB)
        projects = request.POST.get("projects", "")
        previous_company = request.POST.get("previous_company", "")

        if form.is_valid():
            data = form.cleaned_data

            # Multi-select skills list
            skills_list = (
                data["skills"].split(",")
                if isinstance(data["skills"], str)
                else [data["skills"]]
            )

            # Save everything into session
            request.session["resume_data"] = {
                "full_name": data["full_name"],
                "email": data["email"],
                "phone": data["phone"],
                "job_role": data["job_role"],
                "education": data["education"],
                "experience": data["experience"],
                "skills": skills_list,
                "projects": projects,
                "previous_company": previous_company,
            }

            # Redirect (so preview always gets correct data)
            return redirect("resume_preview")

    else:
        form = ResumeForm()

    return render(request, "resume_app/create_resume.html", {"form": form})

# ---------------- RESUME PREVIEW ----------------
def resume_preview(request):
    resume_data = request.session.get("resume_data")

    if not resume_data:
        return HttpResponse("No resume data found. Please generate resume first.")

    return render(request, "resume_app/resume_preview.html", {
        "resume": resume_data
    })

# ---------------- EDIT RESUME ----------------
def edit_resume(request):
    enhanced_resume = None

    if request.method == "POST" and request.FILES.get("resume_file"):
        file = request.FILES["resume_file"]

        # Extract text
        if file.name.endswith(".pdf"):
            text = extract_text(BytesIO(file.read()))
        elif file.name.endswith(".docx"):
            doc = Document(file)
            text = "\n".join([p.text for p in doc.paragraphs])
        else:
            text = file.read().decode("utf-8", errors="ignore")

        # Enhance text
        lines = text.splitlines()
        enhanced_lines = []

        for line in lines:
            if any(k in line.lower() for k in ["education", "experience", "skills"]):
                enhanced_lines.append(line.upper())
            elif "," in line:
                bullets = [f"- {s.strip()}" for s in line.split(",")]
                enhanced_lines.extend(bullets)
            else:
                enhanced_lines.append(line)

        enhanced_resume = "\n".join(enhanced_lines)

    return render(request, "resume_app/edit_resume.html", {
        "enhanced_resume": enhanced_resume
    })


# ---------------- PDF DOWNLOAD ----------------
# PDF generator (ReportLab Platypus)
def download_resume_pdf(request):
    """
    Generates a clean, professional single-column resume PDF
    with blue accent bars and spacing.
    """

    resume_data = request.session.get("resume_data")

    # Fallback for edit-resume text
    if not resume_data:
        resume_text = request.POST.get("resume_text", "")
        if not resume_text:
            return HttpResponse("No resume data provided.", status=400)
        resume_data = {
            "full_name": "Resume",
            "email": "",
            "phone": "",
            "job_role": "",
            "education": "",
            "experience": resume_text,
            "skills": [],
            "projects": "",
            "previous_company": ""
        }

    buffer = BytesIO()

    # PDF document setup
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=50, rightMargin=50,
        topMargin=50, bottomMargin=40
    )

    styles = getSampleStyleSheet()

    # Main title style
    title_style = ParagraphStyle(
        "TitleStyle",
        fontSize=26,
        leading=30,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#003b88"),
        spaceAfter=15,
        bold=True,
    )

    # Section title
    section_title = ParagraphStyle(
        "SectionTitle",
        parent=styles["Heading2"],
        fontSize=15,
        textColor=colors.HexColor("#003b88"),
        spaceBefore=10,
        spaceAfter=4,
    )

    # Normal text
    normal = ParagraphStyle(
        "Normal",
        parent=styles["Normal"],
        fontSize=11,
        leading=15,
    )

    # Small centered line below title
    def separator():
        return Table(
            [[" "]],
            colWidths=[460],
            style=[("LINEBELOW", (0, 0), (-1, -1), 1, colors.HexColor("#003b88"))]
        )

    elements = []

    # ---------------- TITLE ----------------
    elements.append(Paragraph(resume_data["full_name"], title_style))
    elements.append(separator())
    elements.append(Spacer(1, 12))

    # ---------------- CONTACT ----------------
    contact = f"{resume_data.get('email', '')}  |  {resume_data.get('phone', '')}"
    if contact.strip():
        elements.append(Paragraph(contact, normal))
        elements.append(Spacer(1, 15))

    # ---------------- JOB ROLE ----------------
    if resume_data.get("job_role"):
        elements.append(Paragraph("Professional Title", section_title))
        elements.append(Paragraph(resume_data["job_role"], normal))
        elements.append(Spacer(1, 10))

    # ---------------- EDUCATION ----------------
    if resume_data.get("education"):
        elements.append(Paragraph("Education", section_title))
        for line in resume_data["education"].splitlines():
            elements.append(Paragraph("• " + line, normal))
        elements.append(Spacer(1, 10))

    # ---------------- EXPERIENCE ----------------
    if resume_data.get("experience"):
        elements.append(Paragraph("Experience", section_title))
        for line in resume_data["experience"].splitlines():
            elements.append(Paragraph("• " + line, normal))
        elements.append(Spacer(1, 10))

    # ---------------- SKILLS ----------------
    skills = resume_data.get("skills", [])
    if not isinstance(skills, list):
        skills = skills.split(",")

    if skills:
        elements.append(Paragraph("Skills", section_title))
        elements.append(Paragraph(", ".join(skills), normal))
        elements.append(Spacer(1, 10))

    # ---------------- PROJECTS ----------------
    if resume_data.get("projects"):
        elements.append(Paragraph("Projects", section_title))
        for line in resume_data["projects"].splitlines():
            elements.append(Paragraph("• " + line, normal))
        elements.append(Spacer(1, 10))

    # ---------------- PREVIOUS COMPANY ----------------
    if resume_data.get("previous_company"):
        elements.append(Paragraph("Previous Company Details", section_title))
        for line in resume_data["previous_company"].splitlines():
            elements.append(Paragraph("• " + line, normal))
        elements.append(Spacer(1, 10))

    # BORDER ON ALL PAGES
    def add_border(canvas, doc):
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor("#003b88"))
        canvas.setLineWidth(2)
        canvas.roundRect(30, 30, 550, 735, radius=12)
        canvas.restoreState()

    doc.build(elements, onFirstPage=add_border, onLaterPages=add_border)

    buffer.seek(0)
    response = HttpResponse(buffer, content_type="application/pdf")
    response["Content-Disposition"] = 'attachment; filename="resume.pdf"'
    return response

# ---------------- DOCX DOWNLOAD ----------------



def download_resume_doc(request):
    """
    Clean single-column DOCX resume with blue border and proper spacing.
    """

    resume_data = request.session.get("resume_data")

    # Fallback
    if not resume_data:
        txt = request.POST.get("resume_text", "")
        resume_data = {
            "full_name": "Resume",
            "email": "",
            "phone": "",
            "job_role": "",
            "education": "",
            "experience": txt,
            "skills": [],
            "projects": "",
            "previous_company": ""
        }

    doc = DocxDocument()

    # -------- BORDER (FIXED ONLY HERE) --------
    section = doc.sections[0]
    border_xml = f'''
        <w:pgBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="12" w:color="003b88"/>
            <w:left w:val="single" w:sz="12" w:color="003b88"/>
            <w:right w:val="single" w:sz="12" w:color="003b88"/>
            <w:bottom w:val="single" w:sz="12" w:color="003b88"/>
        </w:pgBorders>
    '''
    section._sectPr.append(parse_xml(border_xml))

    # -------- TITLE --------
    title = doc.add_paragraph()
    run = title.add_run(resume_data["full_name"])
    run.bold = True
    run.font.size = Pt(24)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # -------- CONTACT --------
    contact = f"{resume_data.get('email', '')} | {resume_data.get('phone', '')}"
    p = doc.add_paragraph(contact)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()

    # Helper to insert sections
    def add_section(title, content, bullet=False):
        doc.add_heading(title, level=2)
        for line in content:
            if bullet:
                doc.add_paragraph(line, style="List Bullet")
            else:
                doc.add_paragraph(line)

    # JOB ROLE
    if resume_data.get("job_role"):
        doc.add_heading("Professional Title", level=2)
        doc.add_paragraph(resume_data["job_role"])

    # EDUCATION
    if resume_data.get("education"):
        add_section("Education", resume_data["education"].splitlines(), True)

    # EXPERIENCE
    if resume_data.get("experience"):
        add_section("Experience", resume_data["experience"].splitlines(), True)

    # SKILLS
    skills = resume_data.get("skills", [])
    if not isinstance(skills, list):
        skills = skills.split(",")
    if skills:
        doc.add_heading("Skills", level=2)
        doc.add_paragraph(", ".join(skills))

    # PROJECTS
    if resume_data.get("projects"):
        add_section("Projects", resume_data["projects"].splitlines(), True)

    # PREVIOUS COMPANY
    if resume_data.get("previous_company"):
        add_section(
            "Previous Company Details",
            resume_data["previous_company"].splitlines()
        )

    # Export file
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = 'attachment; filename="resume.docx"'
    return response


# ---------------- SKILL GAP ANALYSIS RESULT ----------------
nlp = spacy.load("en_core_web_sm")

@csrf_exempt
def enhanced_skill_gap_analysis(request):
    # ---------------- GET → Show form ----------------
    if request.method == "GET":
        return redirect("skill_gap_form")

    job_role = ""
    found_skills = []
    missing_skills = []
    suggestions = []
    resume_score = 0
    interview_questions = []
    ats_result = {"found": [], "missing": []}

    # ---------------- POST ----------------
    if request.method == "POST" and request.FILES.get("resume_file"):
        file = request.FILES["resume_file"]
        job_role = request.POST.get("job_role", "")
        additional_skills = request.POST.getlist("skills")

        # Extract resume text
        if file.name.endswith(".pdf"):
            text = extract_text(BytesIO(file.read()))
        elif file.name.endswith(".docx"):
            doc = Document(file)
            text = "\n".join([p.text for p in doc.paragraphs])
        else:
            text = file.read().decode("utf-8", errors="ignore")

        text_lower = text.lower()

        # ---------------- SKILL MAP ----------------
        job_skills = {
                "java developer": [
                    "java", "sql", "cloud computing"
                ],

                "python developer": [
                    "python", "sql", "machine learning", "ai"
                ],

                "data analyst": [
                    "python", "sql", "big data", "statistics"
                ],

                "project manager": [
                    "cloud computing", "networking"
                ],

                "data scientist": [
                    "python", "machine learning", "deep learning", "ai", "big data"
                ],

                "aiml engineer": [
                    "python", "machine learning", "deep learning", "ai"
                ],

                "software developer": [
                    "java", "python", "sql"
                ],

                "full stack developer": [
                    "react", "nodejs", "sql", "cloud computing"
                ],

                "frontend developer": [
                    "react", "javascript"
                ],

                "backend developer": [
                    "java", "python", "sql", "nodejs"
                ],

                "cloud engineer": [
                    "cloud computing", "kubernetes", "networking"
                ],

                "devops engineer": [
                    "cloud computing", "kubernetes", "networking"
                ],

                "cybersecurity analyst": [
                    "cybersecurity", "networking", "cloud computing"
                ],

                "big data engineer": [
                    "big data", "python", "sql", "cloud computing"
                ],

                "network engineer": [
                    "networking", "cybersecurity"
                ],

                "ai engineer": [
                    "ai", "machine learning", "deep learning", "python"
                ]
            }


        expected = job_skills.get(job_role.lower(), [])

        for s in additional_skills:
            expected.append(s.lower())

        expected = list(set(expected))

        # ---------------- SKILL MATCH ----------------
        found_skills = [s for s in expected if s in text_lower]
        missing_skills = [s for s in expected if s not in found_skills]

        ats_result = {
            "found": found_skills,
            "missing": missing_skills
        }

        # ---------------- COURSE SUGGESTIONS ----------------
        for skill in missing_skills:
            suggestions.append({
                "skill": skill.title(),
                "courses": [
                    {"title": f"{skill.title()} – Coursera", "link": f"https://www.coursera.org/search?query={skill}"},
                    {"title": f"{skill.title()} – Udemy", "link": f"https://www.udemy.com/courses/search/?q={skill}"},
                    {"title": f"{skill.title()} – YouTube", "link": f"https://www.youtube.com/results?search_query={skill}"}
                ]
            })

        # ---------------- SCORE ----------------
        if expected:
            resume_score = int((len(found_skills) / len(expected)) * 100)

        # ---------------- INTERVIEW QUESTIONS ----------------
        interview_bank = {
        "java developer": [
        "Explain OOP concepts",
        "What is JVM and JDK?",
        "Explain Collections Framework",
        "Difference between abstract class and interface",
        "What is multithreading?",
        "Explain garbage collection",
        "What is exception handling?",
        "Difference between HashMap and Hashtable",
        "What are streams in Java?",
        "Explain SOLID principles"
    ],

    "python developer": [
        "Explain GIL",
        "What are decorators?",
        "Difference between list and tuple",
        "What is list comprehension?",
        "Explain generators",
        "What is lambda function?",
        "Difference between deep copy and shallow copy",
        "What are Python modules and packages?",
        "Explain exception handling in Python",
        "What is virtual environment?"
    ],

    "sql developer": [
        "What is normalization?",
        "Explain primary key and foreign key",
        "Difference between WHERE and HAVING",
        "What are joins?",
        "Explain index in SQL",
        "Difference between DELETE and TRUNCATE",
        "What is a subquery?",
        "Explain ACID properties",
        "What is a stored procedure?",
        "What is view in SQL?"
    ],

    "data analyst": [
        "Explain ETL process",
        "What is A/B testing?",
        "Difference between OLAP and OLTP",
        "What is data cleaning?",
        "Explain descriptive vs inferential statistics",
        "What are KPIs?",
        "Explain data visualization best practices",
        "Difference between mean and median",
        "What is correlation?",
        "Explain pivot tables"
    ],

    "data scientist": [
        "Difference between supervised and unsupervised learning",
        "Explain overfitting and underfitting",
        "What is feature engineering?",
        "Explain bias-variance tradeoff",
        "What is cross validation?",
        "Difference between classification and regression",
        "Explain confusion matrix",
        "What is ROC curve?",
        "What is dimensionality reduction?",
        "Explain PCA"
    ],

    "aiml engineer": [
        "What is machine learning?",
        "Difference between AI and ML",
        "Explain neural networks",
        "What is backpropagation?",
        "Difference between CNN and RNN",
        "Explain activation functions",
        "What is model deployment?",
        "Explain hyperparameter tuning",
        "What is transfer learning?",
        "Explain reinforcement learning"
    ],

    "business analyst": [
        "What is requirement gathering?",
        "Explain BRD and FRD",
        "What are use cases?",
        "Explain SWOT analysis",
        "What is gap analysis?",
        "Difference between business analyst and data analyst",
        "What is stakeholder management?",
        "Explain UML diagrams",
        "What are KPIs in business?",
        "What is risk analysis?"
    ],

    "software developer": [
        "Explain SDLC",
        "What is version control?",
        "Difference between Agile and Waterfall",
        "What is REST API?",
        "Explain MVC architecture",
        "What is debugging?",
        "Explain code optimization",
        "Difference between frontend and backend",
        "What is unit testing?",
        "Explain design patterns"
    ],

    "cloud engineer": [
        "What is cloud computing?",
        "Difference between IaaS, PaaS, SaaS",
        "What is virtualization?",
        "Explain public vs private cloud",
        "What is load balancing?",
        "What is auto scaling?",
        "Explain cloud security",
        "What is containerization?",
        "Difference between Docker and VM",
        "Explain AWS EC2"
    ],

    "cybersecurity analyst": [
        "What is cybersecurity?",
        "Explain CIA triad",
        "What is phishing?",
        "Difference between vulnerability and threat",
        "What is firewall?",
        "Explain encryption and decryption",
        "What is malware?",
        "Explain penetration testing",
        "What is IDS and IPS?",
        "What is zero trust security?"
    ],

    "devops engineer": [
        "What is DevOps?",
        "Explain CI/CD pipeline",
        "What is Docker?",
        "What is Kubernetes?",
        "Explain infrastructure as code",
        "Difference between Jenkins and GitHub Actions",
        "What is monitoring?",
        "Explain configuration management",
        "What is blue-green deployment?",
        "Explain version control systems"
    ],

    "ui/ux designer": [
        "Difference between UI and UX",
        "What is wireframing?",
        "Explain prototyping",
        "What are usability principles?",
        "What is user persona?",
        "Explain user journey map",
        "What is accessibility in design?",
        "Explain design thinking",
        "What are heuristics?",
        "Difference between low-fidelity and high-fidelity design"
    ]
}


        interview_questions = interview_bank.get(job_role.lower(), [])

    # ---------------- FINAL RETURN (CRITICAL) ----------------
    return render(
        request,
        "resume_app/enhance_skill_gap_result.html",
        {
            "job_role": job_role,
            "resume_score": resume_score,
            "suggestions": suggestions,
            "interview_questions": interview_questions,
            "ats_result": ats_result,
        }
    )

#---------------- AI CHAT ASSISTANT ----------------
openai.api_key = "YOUR_OPENAI_API_KEY"

@csrf_exempt
def chat_assistant(request):
    """
    AI Chat Assistant for Resume Builder
    """
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            user_message = data.get("message", "")

            if not user_message.strip():
                return JsonResponse(
                    {"reply": "Please enter a valid message."},
                    status=400
                )

            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "You are an AI Resume Assistant. "
                            "Help users with resume writing, career guidance, "
                            "skill suggestions, and interview preparation."
                        )
                    },
                    {
                        "role": "user",
                        "content": user_message
                    }
                ],
                max_tokens=500,
                temperature=0.7
            )

            assistant_reply = response["choices"][0]["message"]["content"]

            return JsonResponse({"reply": assistant_reply})

        except Exception as e:
            return JsonResponse(
                {"error": str(e)},
                status=500
            )

    return JsonResponse(
        {"error": "Invalid request method"},
        status=405
    )




# ================= BACKEND DEMO ENDPOINT =================

from .models import Skill
from .models import JobRoleSkill
from .models import JobRole


# ================= BACKEND SKILL GAP DEMO =================



def backend_skill_gap(request):
    """
    Backend-only Skill Gap Analysis (No frontend)
    Demonstrates how backend processes resume vs job role
    """

    resume_text = request.GET.get("resume_text", "").lower()
    job_role = request.GET.get("job_role", "").lower()

    if not resume_text or not job_role:
        return JsonResponse({
            "error": "resume_text and job_role are required"
        }, status=400)

    # Example job-role → skills mapping (backend logic)
    job_skill_map = {
        "python developer": ["python", "django", "sql", "machine learning"],
        "java developer": ["java", "sql", "cloud computing"],
        "data scientist": ["python", "machine learning", "deep learning", "sql"]
    }

    expected_skills = job_skill_map.get(job_role, [])

    found = []
    missing = []

    for skill in expected_skills:
        if skill in resume_text:
            found.append(skill)
        else:
            missing.append(skill)

    return JsonResponse({
        "feature": "Skill Gap Analysis",
        "job_role": job_role,
        "expected_skills": expected_skills,
        "found_skills": found,
        "missing_skills": missing,
        "status": "Skill gap analysis completed"
    })
# ================= BACKEND DATABASE STATUS =================

def backend_database_status(request):
    """
    Backend-only endpoint to prove database connectivity
    """

    try:
        skill_count = Skill.objects.count()
        role_count = JobRole.objects.count()
        role_skill_count = JobRoleSkill.objects.count()

        return JsonResponse({
            "feature": "Database Connectivity",
            "database": "Connected",
            "tables": {
                "Skill": skill_count,
                "JobRole": role_count,
                "JobRoleSkill": role_skill_count
            },
            "status": "OK"
        })

    except Exception as e:
        return JsonResponse({
            "feature": "Database Connectivity",
            "database": "Error",
            "error": str(e)
        }, status=500)

# ================= BACKEND JD KEYWORD ANALYSIS =================





@csrf_exempt
def backend_jd_keyword_analysis(request):
    """
    Backend-only JD Keyword Analysis
    Compares resume text with job description keywords
    """

    resume_text = request.GET.get("resume_text", "").lower()
    jd_text = request.GET.get("jd_text", "").lower()

    if not resume_text or not jd_text:
        return JsonResponse(
            {"error": "resume_text and jd_text are required"},
            status=400
        )

    # Extract keywords
    resume_keywords = set(re.findall(r'\b[a-zA-Z]{3,}\b', resume_text))
    jd_keywords = set(re.findall(r'\b[a-zA-Z]{3,}\b', jd_text))

    matched_keywords = sorted(resume_keywords & jd_keywords)

    match_percentage = 0
    if jd_keywords:
        match_percentage = int((len(matched_keywords) / len(jd_keywords)) * 100)

    return JsonResponse({
        "feature": "JD Keyword Analysis",
        "resume_keywords": sorted(resume_keywords),
        "jd_keywords": sorted(jd_keywords),
        "matched_keywords": matched_keywords,
        "match_percentage": match_percentage,
        "status": "JD keyword analysis completed"
    })
