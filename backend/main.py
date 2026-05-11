from fastapi import FastAPI, HTTPException, Depends, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import HTTPBearer
from fastapi.responses import StreamingResponse
import jwt
import bcrypt
import os
import datetime
import asyncio
from dotenv import load_dotenv
from typing import Optional
import json
import io
from supabase import create_client, Client
from groq import Groq

load_dotenv(os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env"))

supabase_url = os.getenv("SUPABASE_URL")
supabase_key = os.getenv("SUPABASE_KEY")
app_secret = os.getenv("APP_SECRET_KEY")

if not supabase_url or not supabase_key:
    raise RuntimeError("SUPABASE_URL and SUPABASE_KEY must be set in .env")
if not app_secret:
    raise RuntimeError("APP_SECRET_KEY must be set in .env")

supabase: Client = create_client(supabase_url, supabase_key)
supabase_available = True
print("Supabase connected successfully")

groq_client = Groq(api_key=os.getenv("GROQ_API_KEY"))

# Initialize FastAPI app
app = FastAPI(title="AI Resume Builder API")

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=[os.getenv("FRONTEND_URL", "http://localhost:5173")],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Security
security = HTTPBearer(auto_error=False)

def get_optional_user(credentials=Depends(security)):
    if not credentials:
        return None
    try:
        payload = jwt.decode(
            credentials.credentials,
            app_secret,
            algorithms=["HS256"]
        )
        return payload
    except Exception:
        raise HTTPException(401, "Invalid token")

def require_user(user=Depends(get_optional_user)):
    if not user:
        raise HTTPException(401, "Login required")
    return user

# Import our modules
from skill_gap import analyze_skill_gap
from jd_matcher import analyze_jd
from resume_engine import generate_pdf, generate_docx
from llm_service import generate_learning_roadmap, generate_cover_letter, generate_ats_score

# Routes
@app.get("/health")
async def health_check():
    return {"status": "ok"}

@app.post("/api/auth/signup")
async def signup(user_data: dict):
    if not supabase_available:
        raise HTTPException(503, "Authentication service not configured")

    try:
        # Check if userid already exists
        existing_userid = supabase.table("users").select("userid").eq("userid", user_data["userid"]).execute()
        if existing_userid.data:
            raise HTTPException(400, "User ID already exists")

        # Check if email already exists
        existing_email = supabase.table("users").select("email").eq("email", user_data["email"]).execute()
        if existing_email.data:
            raise HTTPException(400, "Email already exists")

        # Hash the password with bcrypt before storing
        hashed_pw = bcrypt.hashpw(
            user_data["password"].encode("utf-8"),
            bcrypt.gensalt()
        ).decode("utf-8")

        # Insert new user with hashed password
        supabase.table("users").insert({
            "name": user_data["name"],
            "email": user_data["email"],
            "userid": user_data["userid"],
            "phone": user_data.get("phone"),
            "password": hashed_pw
        }).execute()

        return {"message": "Account created successfully"}
    except Exception as e:
        if isinstance(e, HTTPException):
            raise e
        raise HTTPException(500, "Internal server error")

@app.post("/api/auth/login")
async def login(credentials: dict):
    if not supabase_available:
        raise HTTPException(503, "Authentication service not configured")

    try:
        # Fetch user by userid only — never query by plaintext password
        result = supabase.table("users").select("*").eq("userid", credentials["userid"]).execute()

        if not result.data:
            raise HTTPException(401, "Invalid credentials")

        user = result.data[0]

        # Verify the submitted password against the stored bcrypt hash
        if not bcrypt.checkpw(
            credentials["password"].encode("utf-8"),
            user["password"].encode("utf-8")
        ):
            raise HTTPException(401, "Invalid credentials")

        # Issue a signed JWT valid for 7 days
        payload = {
            "sub": user["id"],          # Supabase UUID — used by protected routes
            "userid": user["userid"],
            "name": user["name"],
            "email": user["email"],
            "exp": datetime.datetime.utcnow() + datetime.timedelta(days=7)
        }
        token = jwt.encode(payload, app_secret, algorithm="HS256")

        return {
            "message": "Login successful",
            "token": token,
            "user": {
                "userid": user["userid"],
                "name": user["name"],
                "email": user["email"]
            }
        }
    except Exception as e:
        if isinstance(e, HTTPException):
            raise e
        raise HTTPException(500, "Internal server error")

def _fetch_user_resumes(user: dict):
    user_id = user.get("sub")
    if not supabase_available or not user_id:
        return {"resumes": []}
    try:
        result = (
            supabase.table("resumes")
            .select("id, full_name, job_title, created_at")
            .eq("user_id", user_id)
            .order("created_at", desc=True)
            .execute()
        )
        return {"resumes": result.data or []}
    except Exception as e:
        print(f"List resumes: {e}")
        return {"resumes": []}


@app.post("/api/resume/save")
async def save_resume(data: dict):
    if not supabase_available:
        return {"ok": False}
    try:
        userid = data.get("userid")
        if not userid:
            return {"ok": False}
        user = supabase.table("users").select("id").eq("userid", userid).execute()
        if not user.data:
            return {"ok": False}
        user_id = user.data[0]["id"]
        supabase.table("resumes").insert({
            "user_id": user_id,
            "full_name": data.get("full_name"),
            "job_title": data.get("job_title"),
            "resume_data": data.get("resume_data")
        }).execute()
        return {"ok": True}
    except Exception as e:
        print(f"Save resume error: {e}")
        return {"ok": False}


@app.post("/api/resume/create")
async def create_resume(
    full_name: str = Form(...),
    email: str = Form(...),
    phone: str = Form(...),
    job_role: str = Form(...),
    education: str = Form(...),
    experience: str = Form(...),
    skills: str = Form(...),
    projects: str = Form(...),
    previous_company: str = Form(...),
    user=Depends(require_user)
):
    resume_data = {
        "full_name": full_name,
        "email": email,
        "phone": phone,
        "job_role": job_role,
        "education": education,
        "experience": experience,
        "skills": skills.split(",") if skills else [],
        "projects": projects,
        "previous_company": previous_company
    }
    saved_id = None
    uid = user.get("sub")
    if supabase_available and uid:
        try:
            row = {
                "user_id": uid,
                "full_name": full_name,
                "job_title": job_role,
                "resume_data": resume_data,
            }
            ins = supabase.table("resumes").insert(row).execute()
            if ins.data:
                saved_id = ins.data[0].get("id")
        except Exception as e:
            print(f"Resume persist skipped: {e}")
    return {"message": "Resume created successfully", "data": resume_data, "id": saved_id}


@app.get("/api/resume/{resume_id}/pdf")
async def download_saved_resume_pdf(resume_id: str, user=Depends(require_user)):
    user_id = user.get("sub")
    if not supabase_available or not user_id:
        raise HTTPException(status_code=503, detail="Resume storage is not configured")

    try:
        rows = (
            supabase.table("resumes")
            .select("resume_data")
            .eq("id", resume_id)
            .eq("user_id", user_id)
            .execute()
        )
        if not rows.data:
            raise HTTPException(status_code=404, detail="Resume not found")
        raw = rows.data[0]["resume_data"]
        resume_blob = json.loads(raw) if isinstance(raw, str) else raw
        pdf_bytes = generate_pdf(resume_blob)
        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="resume_{resume_id}.pdf"'},
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/resume/{resume_id}")
async def delete_saved_resume(resume_id: str, userid: str = None):
    if not supabase_available or not userid:
        raise HTTPException(status_code=400, detail="userid required")
    try:
        user = supabase.table("users").select("id").eq("userid", userid).execute()
        if not user.data:
            raise HTTPException(status_code=404, detail="User not found")
        user_id = user.data[0]["id"]
        supabase.table("resumes").delete().eq("id", resume_id).eq("user_id", user_id).execute()
        return {"ok": True}
    except HTTPException:
        raise
    except Exception as e:
        print(f"Delete resume error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/resume/download-pdf")
async def download_resume_pdf(resume_data: dict):
    pdf_bytes = generate_pdf(resume_data)

    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=resume.pdf"}
    )

@app.post("/api/resume/download-docx")
async def download_resume_docx(resume_data: dict):
    docx_bytes = generate_docx(resume_data)

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=resume.docx"}
    )

@app.post("/api/resume/ats-score")
async def ats_score_analyze(resume_data: dict):
    result = generate_ats_score(resume_data)
    return result

@app.post("/api/skill-gap/analyze")
async def skill_gap_analyze(
    job_role: str = Form(...),
    resume_file: UploadFile = File(...),
    additional_skills: str = Form("")
):
    # Validate file
    if not resume_file.filename.lower().endswith(('.pdf', '.docx')):
        raise HTTPException(400, "Only PDF and DOCX files are allowed")

    # Extract text
    content = await resume_file.read()
    if resume_file.filename.lower().endswith('.pdf'):
        from pdfminer.high_level import extract_text
        resume_text = extract_text(io.BytesIO(content))
    else:
        from docx import Document
        doc = Document(io.BytesIO(content))
        resume_text = "\n".join([p.text for p in doc.paragraphs])

    result = analyze_skill_gap(job_role, resume_text, additional_skills)
    if result["ats_result"]["missing"]:
        try:
            print(f"Generating roadmap for {job_role} with missing skills: {result['ats_result']['missing']}")
            roadmap = await asyncio.get_event_loop().run_in_executor(
                None,
                lambda: generate_learning_roadmap(job_role, result["ats_result"]["missing"])
            )
            result["roadmap"] = roadmap
        except:
            result["roadmap"] = None
    
    # Generate interview questions
    try:
        iq_prompt = f"""Generate 5 interview questions with detailed answers for a {job_role} role.
    Return ONLY valid JSON in this exact format, no markdown:
    [
        {{"question": "Question here?", "answer": "Detailed answer here."}}
    ]"""
        iq_response = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": iq_prompt}],
            max_tokens=1500
        )
        iq_content = iq_response.choices[0].message.content.strip()
        if iq_content.startswith("```json"): iq_content = iq_content[7:]
        if iq_content.startswith("```"): iq_content = iq_content[3:]
        if iq_content.endswith("```"): iq_content = iq_content[:-3]
        result["interview_questions"] = json.loads(iq_content.strip())
    except Exception as e:
        print(f"Interview Q gen error: {e}")
        result["interview_questions"] = []
    
    return result

@app.post("/api/jd-match")
async def jd_match_analyze(
    job_description: str = Form(...),
    resume: UploadFile = File(...)
):
    if not job_description.strip():
        raise HTTPException(400, "Job description is required")
    if not resume.filename.lower().endswith(('.pdf', '.docx')):
        raise HTTPException(400, "Only PDF and DOCX files are allowed")

    content = await resume.read()
    if resume.filename.lower().endswith('.pdf'):
        from pdfminer.high_level import extract_text
        resume_text = extract_text(io.BytesIO(content))
    else:
        from docx import Document
        doc = Document(io.BytesIO(content))
        resume_text = "\n".join([p.text for p in doc.paragraphs])

    result = analyze_jd(job_description, resume_text)
    return result

@app.post("/api/jd-match-upload")
async def jd_match_upload(resume_file: UploadFile = File(...)):
    if not resume_file.filename.lower().endswith(('.pdf', '.docx')):
        raise HTTPException(400, "Only PDF and DOCX files are allowed")

    content = await resume_file.read()
    if resume_file.filename.lower().endswith('.pdf'):
        from pdfminer.high_level import extract_text
        resume_text = extract_text(io.BytesIO(content))
    else:
        from docx import Document
        doc = Document(io.BytesIO(content))
        resume_text = "\n".join([p.text for p in doc.paragraphs])

    result = analyze_jd(resume_text, resume_text)
    return result

@app.post("/api/cover-letter/generate")
async def cover_letter_generate(data: dict):
    try:
        # Validate required fields
        required = ["full_name", "job_title", "skills", "work_experience"]
        for field in required:
            if field not in data:
                raise HTTPException(400, f"Field '{field}' is required")
        
        cover_letter = generate_cover_letter(data)
        return {"cover_letter": cover_letter}
    except HTTPException as e:
        raise e
    except Exception as e:
        print(f"Cover letter generation error: {e}")
        raise HTTPException(500, "Failed to generate cover letter")

@app.get("/api/my-resumes")
async def list_my_resumes(userid: Optional[str] = None):
    if not userid or not supabase_available:
        return {"resumes": []}
    try:
        user = supabase.table("users").select("id").eq("userid", userid).execute()
        if not user.data:
            return {"resumes": []}
        user_id = user.data[0]["id"]
        result = (
            supabase.table("resumes")
            .select("id, full_name, job_title, created_at")
            .eq("user_id", user_id)
            .order("created_at", desc=True)
            .execute()
        )
        return {"resumes": result.data or []}
    except Exception as e:
        print(f"List resumes error: {e}")
        return {"resumes": []}


@app.get("/api/resumes/me")
async def get_user_resumes(user=Depends(require_user)):
    return _fetch_user_resumes(user)